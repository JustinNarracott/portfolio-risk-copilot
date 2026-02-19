"""Unit tests for DOCX generator (Issues #18, #19, #20, #22)."""

from datetime import date
from pathlib import Path

import pytest
from docx import Document

from src.ingestion.parser import parse_file
from src.risk_engine.engine import (
    PortfolioRiskReport,
    ProjectRiskSummary,
    Risk,
    RiskCategory,
    RiskSeverity,
    analyse_portfolio,
)
from src.artefacts.docx_generator import (
    BrandConfig,
    generate_board_briefing,
    generate_steering_pack,
    generate_project_status_pack,
)

SAMPLE_CSV = Path(__file__).parent.parent.parent / "sample-data" / "jira-export-sample.csv"
REF_DATE = date(2026, 2, 19)


@pytest.fixture()
def report() -> PortfolioRiskReport:
    projects = parse_file(SAMPLE_CSV)
    return analyse_portfolio(projects, top_n=5, reference_date=REF_DATE)


# ──────────────────────────────────────────────
# Board briefing (Issue #18)
# ──────────────────────────────────────────────


class TestBoardBriefing:

    def test_generates_valid_docx(self, report, tmp_path):
        out = tmp_path / "board.docx"
        result = generate_board_briefing(report, output_path=out)
        assert result.exists()
        doc = Document(str(result))
        assert len(doc.paragraphs) > 0

    def test_contains_portfolio_status(self, report, tmp_path):
        out = tmp_path / "board.docx"
        generate_board_briefing(report, output_path=out)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Portfolio" in full_text

    def test_contains_rag_status(self, report, tmp_path):
        out = tmp_path / "board.docx"
        generate_board_briefing(report, output_path=out)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert report.portfolio_rag in full_text

    def test_contains_risk_section(self, report, tmp_path):
        out = tmp_path / "board.docx"
        generate_board_briefing(report, output_path=out)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Risk" in full_text

    def test_contains_decisions(self, report, tmp_path):
        out = tmp_path / "board.docx"
        generate_board_briefing(report, output_path=out)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Recommended" in full_text or "Decision" in full_text or "decision" in full_text

    def test_has_table(self, report, tmp_path):
        out = tmp_path / "board.docx"
        generate_board_briefing(report, output_path=out)
        doc = Document(str(out))
        assert len(doc.tables) >= 1

    def test_table_has_all_projects(self, report, tmp_path):
        out = tmp_path / "board.docx"
        generate_board_briefing(report, output_path=out)
        doc = Document(str(out))
        table = doc.tables[0]
        table_text = " ".join(cell.text for row in table.rows for cell in row.cells)
        for summary in report.project_summaries:
            assert summary.project_name in table_text

    def test_default_output_path(self, report, tmp_path, monkeypatch):
        monkeypatch.chdir(tmp_path)
        result = generate_board_briefing(report)
        assert result.name == "board-briefing.docx"


# ──────────────────────────────────────────────
# Steering committee pack (Issue #19)
# ──────────────────────────────────────────────


class TestSteeringPack:

    def test_generates_valid_docx(self, report, tmp_path):
        out = tmp_path / "steering.docx"
        result = generate_steering_pack(report, output_path=out)
        assert result.exists()
        doc = Document(str(result))
        assert len(doc.paragraphs) > 10  # Should be substantial

    def test_contains_exec_summary(self, report, tmp_path):
        out = tmp_path / "steering.docx"
        generate_steering_pack(report, output_path=out)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Executive Summary" in full_text or "Exec" in full_text

    def test_contains_top_5_risks(self, report, tmp_path):
        out = tmp_path / "steering.docx"
        generate_steering_pack(report, output_path=out)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        # Should have at least some severity indicators
        assert "[Critical]" in full_text or "[High]" in full_text

    def test_contains_talking_points(self, report, tmp_path):
        out = tmp_path / "steering.docx"
        generate_steering_pack(report, output_path=out)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Talking" in full_text or "talking" in full_text

    def test_has_detailed_table(self, report, tmp_path):
        out = tmp_path / "steering.docx"
        generate_steering_pack(report, output_path=out)
        doc = Document(str(out))
        assert len(doc.tables) >= 1
        # Detailed table should have 5 columns
        table = doc.tables[0]
        assert len(table.columns) == 5

    def test_more_content_than_board_briefing(self, report, tmp_path):
        board_path = tmp_path / "board.docx"
        steering_path = tmp_path / "steering.docx"
        generate_board_briefing(report, output_path=board_path)
        generate_steering_pack(report, output_path=steering_path)
        # Steering should have more paragraphs
        board_doc = Document(str(board_path))
        steering_doc = Document(str(steering_path))
        assert len(steering_doc.paragraphs) > len(board_doc.paragraphs)


# ──────────────────────────────────────────────
# Project status pack (Issue #20)
# ──────────────────────────────────────────────


class TestProjectStatusPack:

    def test_generates_valid_docx(self, report, tmp_path):
        out = tmp_path / "status.docx"
        result = generate_project_status_pack(report, output_path=out)
        assert result.exists()

    def test_contains_all_projects(self, report, tmp_path):
        out = tmp_path / "status.docx"
        generate_project_status_pack(report, output_path=out)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        for summary in report.project_summaries:
            assert summary.project_name in full_text

    def test_contains_rag_per_project(self, report, tmp_path):
        out = tmp_path / "status.docx"
        generate_project_status_pack(report, output_path=out)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Status:" in full_text

    def test_contains_action_items(self, report, tmp_path):
        out = tmp_path / "status.docx"
        generate_project_status_pack(report, output_path=out)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Action" in full_text


# ──────────────────────────────────────────────
# Template customisation (Issue #22)
# ──────────────────────────────────────────────


class TestBrandCustomisation:

    def test_custom_colours_applied(self, report, tmp_path):
        brand = BrandConfig(primary_colour="990000", accent_colour="CC3333")
        out = tmp_path / "custom.docx"
        generate_board_briefing(report, brand=brand, output_path=out)
        assert out.exists()

    def test_custom_headings(self, report, tmp_path):
        brand = BrandConfig(custom_headings={
            "board_title": "Monthly Portfolio Update",
            "top_risks": "Key Risk Areas",
        })
        out = tmp_path / "custom.docx"
        generate_board_briefing(report, brand=brand, output_path=out)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Monthly Portfolio Update" in full_text
        assert "Key Risk Areas" in full_text

    def test_custom_font(self, report, tmp_path):
        brand = BrandConfig(heading_font="Arial", body_font="Arial")
        out = tmp_path / "custom.docx"
        generate_board_briefing(report, brand=brand, output_path=out)
        assert out.exists()

    def test_logo_with_valid_path(self, report, tmp_path):
        # Create a tiny valid PNG
        logo_path = tmp_path / "logo.png"
        _create_tiny_png(logo_path)

        brand = BrandConfig(logo_path=str(logo_path))
        out = tmp_path / "logo_test.docx"
        generate_board_briefing(report, brand=brand, output_path=out)
        assert out.exists()

    def test_logo_with_invalid_path_ignored(self, report, tmp_path):
        brand = BrandConfig(logo_path="/nonexistent/logo.png")
        out = tmp_path / "no_logo.docx"
        generate_board_briefing(report, brand=brand, output_path=out)
        assert out.exists()

    def test_customisation_applies_to_steering(self, report, tmp_path):
        brand = BrandConfig(custom_headings={
            "steering_title": "Custom Steering Title",
        })
        out = tmp_path / "steering_custom.docx"
        generate_steering_pack(report, brand=brand, output_path=out)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Custom Steering Title" in full_text

    def test_customisation_applies_to_status_pack(self, report, tmp_path):
        brand = BrandConfig(company_name="Acme Corp")
        out = tmp_path / "status_custom.docx"
        generate_project_status_pack(report, brand=brand, output_path=out)
        assert out.exists()


def _create_tiny_png(path: Path) -> None:
    """Create a minimal valid 1x1 PNG file."""
    import struct, zlib
    sig = b"\x89PNG\r\n\x1a\n"
    ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    ihdr_crc = zlib.crc32(b"IHDR" + ihdr_data) & 0xFFFFFFFF
    ihdr = struct.pack(">I", 13) + b"IHDR" + ihdr_data + struct.pack(">I", ihdr_crc)
    raw = zlib.compress(b"\x00\xFF\x00\x00")
    idat_crc = zlib.crc32(b"IDAT" + raw) & 0xFFFFFFFF
    idat = struct.pack(">I", len(raw)) + b"IDAT" + raw + struct.pack(">I", idat_crc)
    iend_crc = zlib.crc32(b"IEND") & 0xFFFFFFFF
    iend = struct.pack(">I", 0) + b"IEND" + struct.pack(">I", iend_crc)
    path.write_bytes(sig + ihdr + idat + iend)
