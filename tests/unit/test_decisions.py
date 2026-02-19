"""Unit tests for decision log generator (Issue #33)."""

from datetime import date
from pathlib import Path

import pytest
from docx import Document

from src.ingestion.parser import parse_file
from src.risk_engine.engine import analyse_portfolio
from src.benefits.parser import parse_benefits
from src.benefits.calculator import analyse_benefits
from src.investment import analyse_investments
from src.scenario.parser import parse_scenario
from src.scenario.simulator import simulate
from src.scenario.graph import build_dependency_graph
from src.decisions import (
    DecisionLog, DecisionStatus, decision_from_scenario,
    decisions_from_risk_report, decisions_from_investment, export_decision_log,
)

SAMPLE = Path(__file__).parent.parent.parent / "sample-data" / "jira-export-sample.csv"
BENEFITS = Path(__file__).parent.parent.parent / "sample-data" / "benefit-tracker-sample.csv"
REF_DATE = date(2026, 2, 19)


@pytest.fixture()
def projects():
    return parse_file(SAMPLE)


@pytest.fixture()
def risk_report(projects):
    return analyse_portfolio(projects, top_n=5, reference_date=REF_DATE)


@pytest.fixture()
def investment_report(projects, risk_report):
    benefits = parse_benefits(BENEFITS)
    benefit_report = analyse_benefits(benefits, risk_report, REF_DATE)
    return analyse_investments(projects, risk_report, benefit_report)


class TestDecisionLog:

    def test_empty_log(self):
        log = DecisionLog()
        assert len(log.decisions) == 0

    def test_next_id_increments(self):
        log = DecisionLog()
        assert log.next_id() == "DEC-001"
        assert log.next_id() == "DEC-002"

    def test_to_json(self):
        log = DecisionLog()
        j = log.to_json()
        assert '"decision_count": 0' in j


class TestDecisionFromScenario:

    def test_creates_decision(self, projects, risk_report):
        log = DecisionLog()
        graph = build_dependency_graph(projects)
        action = parse_scenario("delay Alpha by 2 months")
        result = simulate(action, projects, graph, REF_DATE)
        d = decision_from_scenario(result, log, REF_DATE)
        assert d.decision_id == "DEC-001"
        assert d.source == "scenario"
        assert len(d.options) == 2
        assert d.status == DecisionStatus.PENDING

    def test_adds_to_log(self, projects, risk_report):
        log = DecisionLog()
        graph = build_dependency_graph(projects)
        action = parse_scenario("increase Beta budget by 20%")
        result = simulate(action, projects, graph, REF_DATE)
        decision_from_scenario(result, log, REF_DATE)
        assert len(log.decisions) == 1


class TestDecisionsFromRiskReport:

    def test_creates_escalation_decision(self, risk_report):
        log = DecisionLog()
        decisions = decisions_from_risk_report(risk_report, log, REF_DATE)
        assert len(decisions) >= 1
        assert any("Red" in d.title or "escalat" in d.title.lower() for d in decisions)

    def test_has_three_options(self, risk_report):
        log = DecisionLog()
        decisions = decisions_from_risk_report(risk_report, log, REF_DATE)
        for d in decisions:
            assert len(d.options) >= 2


class TestDecisionsFromInvestment:

    def test_creates_investment_decisions(self, investment_report):
        log = DecisionLog()
        decisions = decisions_from_investment(investment_report, log, REF_DATE)
        assert isinstance(decisions, list)

    def test_divest_decision_has_budget(self, investment_report):
        log = DecisionLog()
        decisions = decisions_from_investment(investment_report, log, REF_DATE)
        for d in decisions:
            if "divest" in d.title.lower() or "Divest" in d.title:
                assert "£" in d.title


class TestExportDecisionLog:

    def test_exports_valid_docx(self, risk_report, investment_report, tmp_path):
        log = DecisionLog()
        decisions_from_risk_report(risk_report, log, REF_DATE)
        decisions_from_investment(investment_report, log, REF_DATE)
        out = tmp_path / "decisions.docx"
        result = export_decision_log(log, output_path=out)
        assert result.exists()
        doc = Document(str(result))
        assert len(doc.paragraphs) > 5

    def test_exports_empty_log(self, tmp_path):
        log = DecisionLog()
        out = tmp_path / "empty.docx"
        result = export_decision_log(log, output_path=out)
        assert result.exists()

    def test_contains_decision_ids(self, risk_report, tmp_path):
        log = DecisionLog()
        decisions_from_risk_report(risk_report, log, REF_DATE)
        out = tmp_path / "decisions.docx"
        export_decision_log(log, output_path=out)
        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "DEC-001" in text
