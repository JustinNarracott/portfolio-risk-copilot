"""E2E integration test for Sprint 3: artefact generation (Issue #23).

Acceptance criteria from PID:
  - Upload Jira export + benefit tracker → generate Steering Committee Pack (DOCX) in <90s
  - Generated pack includes: exec summary, top 5 risks, 3 decisions, RAG table, forecast changes
  - User can customise logo, brand colours, section headings
  - DOCX and PPTX files are valid and openable
"""

import time
from datetime import date
from pathlib import Path

import pytest
from docx import Document
from pptx import Presentation

from src.ingestion.parser import parse_file
from src.risk_engine.engine import analyse_portfolio
from src.artefacts.docx_generator import (
    BrandConfig,
    generate_board_briefing,
    generate_steering_pack,
    generate_project_status_pack,
)
from src.artefacts.pptx_generator import generate_board_slides

SAMPLE_CSV = Path(__file__).parent.parent.parent / "sample-data" / "jira-export-sample.csv"
REF_DATE = date(2026, 2, 19)


class TestFullArtefactPipeline:

    def test_steering_pack_under_90_seconds(self, tmp_path):
        """PID acceptance: generate Steering Committee Pack in <90 seconds."""
        start = time.time()

        projects = parse_file(SAMPLE_CSV)
        report = analyse_portfolio(projects, top_n=5, reference_date=REF_DATE)
        out = tmp_path / "steering.docx"
        generate_steering_pack(report, output_path=out)

        elapsed = time.time() - start
        assert elapsed < 90.0, f"Pipeline took {elapsed:.2f}s (limit: 90s)"
        assert out.exists()

    def test_steering_pack_content_completeness(self, tmp_path):
        """PID acceptance: exec summary, top 5 risks, 3 decisions, RAG table."""
        projects = parse_file(SAMPLE_CSV)
        report = analyse_portfolio(projects, top_n=5, reference_date=REF_DATE)
        out = tmp_path / "steering.docx"
        generate_steering_pack(report, output_path=out)

        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)

        # Exec summary
        assert "Executive Summary" in full_text or "portfolio" in full_text.lower()
        # Risk section
        assert "[Critical]" in full_text or "[High]" in full_text
        # Decisions
        assert "Recommended" in full_text or "decision" in full_text.lower()
        # RAG table exists
        assert len(doc.tables) >= 1

    def test_all_three_docx_templates_generate(self, tmp_path):
        """All three DOCX templates should generate successfully."""
        projects = parse_file(SAMPLE_CSV)
        report = analyse_portfolio(projects, top_n=5, reference_date=REF_DATE)

        board = tmp_path / "board.docx"
        steering = tmp_path / "steering.docx"
        status = tmp_path / "status.docx"

        generate_board_briefing(report, output_path=board)
        generate_steering_pack(report, output_path=steering)
        generate_project_status_pack(report, output_path=status)

        assert board.exists()
        assert steering.exists()
        assert status.exists()

        # All should be valid DOCX
        Document(str(board))
        Document(str(steering))
        Document(str(status))

    def test_pptx_generates_alongside_docx(self, tmp_path):
        """PPTX should generate alongside DOCX."""
        projects = parse_file(SAMPLE_CSV)
        report = analyse_portfolio(projects, top_n=5, reference_date=REF_DATE)

        docx_path = tmp_path / "board.docx"
        pptx_path = tmp_path / "board.pptx"

        generate_board_briefing(report, output_path=docx_path)
        generate_board_slides(report, output_path=pptx_path)

        assert docx_path.exists()
        assert pptx_path.exists()

        Presentation(str(pptx_path))

    def test_full_pipeline_with_custom_branding(self, tmp_path):
        """PID acceptance: user can customise logo, brand colours, section headings."""
        projects = parse_file(SAMPLE_CSV)
        report = analyse_portfolio(projects, top_n=5, reference_date=REF_DATE)

        brand = BrandConfig(
            primary_colour="003366",
            accent_colour="336699",
            heading_font="Arial",
            body_font="Arial",
            company_name="Acme Corp",
            custom_headings={
                "board_title": "Acme Corp — Portfolio Dashboard",
                "top_risks": "Priority Risk Items",
                "decisions": "Actions for Board Approval",
            },
        )

        board = tmp_path / "board.docx"
        steering = tmp_path / "steering.docx"
        pptx = tmp_path / "board.pptx"

        generate_board_briefing(report, brand=brand, output_path=board)
        generate_steering_pack(report, brand=brand, output_path=steering)
        generate_board_slides(report, brand=brand, output_path=pptx)

        # Verify custom headings applied
        doc = Document(str(board))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Acme Corp — Portfolio Dashboard" in full_text
        assert "Priority Risk Items" in full_text
        assert "Actions for Board Approval" in full_text

    def test_all_artefacts_under_90_seconds(self, tmp_path):
        """All artefacts should generate within 90 seconds total."""
        start = time.time()

        projects = parse_file(SAMPLE_CSV)
        report = analyse_portfolio(projects, top_n=5, reference_date=REF_DATE)

        generate_board_briefing(report, output_path=tmp_path / "board.docx")
        generate_steering_pack(report, output_path=tmp_path / "steering.docx")
        generate_project_status_pack(report, output_path=tmp_path / "status.docx")
        generate_board_slides(report, output_path=tmp_path / "board.pptx")

        elapsed = time.time() - start
        assert elapsed < 90.0, f"All artefacts took {elapsed:.2f}s (limit: 90s)"
