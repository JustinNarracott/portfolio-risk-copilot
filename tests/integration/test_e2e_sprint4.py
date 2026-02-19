"""E2E integration test for Sprint 4: full workflow (Issue #28).

Acceptance criteria from PID:
  - Plugin installs and full workflow runs: ingest → risks → brief steering in <120s
  - Cross-format consistency: CSV, JSON, Excel produce same results
  - All CLI commands functional
  - All artefacts generated and valid
"""

import json
import time
from datetime import date
from pathlib import Path

import pytest
from docx import Document
from pptx import Presentation

from src.cli import main, _session

SAMPLE_DIR = Path(__file__).parent.parent.parent / "sample-data"


@pytest.fixture(autouse=True)
def reset_session():
    _session.projects = []
    _session.report = None
    _session.graph = None
    _session.brand.__init__()
    _session.reference_date = date(2026, 2, 19)
    yield


class TestFullWorkflowE2E:

    def test_full_workflow_under_120_seconds(self, tmp_path, capsys):
        """PID acceptance: ingest → risks → brief steering in <120s."""
        start = time.time()

        assert main(["ingest", str(SAMPLE_DIR)]) == 0
        assert main(["risks"]) == 0
        assert main(["brief", "steering", "--output-dir", str(tmp_path)]) == 0

        elapsed = time.time() - start
        assert elapsed < 120.0, f"Full workflow took {elapsed:.2f}s (limit: 120s)"
        assert (tmp_path / "steering-committee-pack.docx").exists()

    def test_full_workflow_all_artefacts(self, tmp_path, capsys):
        """Generate all artefact types in one workflow."""
        assert main(["ingest", str(SAMPLE_DIR)]) == 0
        assert main(["brief", "all", "--output-dir", str(tmp_path)]) == 0

        assert (tmp_path / "board-briefing.docx").exists()
        assert (tmp_path / "board-briefing.pptx").exists()
        assert (tmp_path / "steering-committee-pack.docx").exists()
        assert (tmp_path / "project-status-pack.docx").exists()

        # Validate all files are openable
        Document(str(tmp_path / "board-briefing.docx"))
        Document(str(tmp_path / "steering-committee-pack.docx"))
        Document(str(tmp_path / "project-status-pack.docx"))
        Presentation(str(tmp_path / "board-briefing.pptx"))

    def test_scenario_then_brief(self, tmp_path, capsys):
        """Run scenario simulation then generate briefing."""
        assert main(["ingest", str(SAMPLE_DIR)]) == 0
        assert main(["scenario", "cut Beta scope by 30%"]) == 0
        assert main(["brief", "board", "--output-dir", str(tmp_path)]) == 0
        assert (tmp_path / "board-briefing.docx").exists()

    def test_risks_json_pipeline(self, tmp_path, capsys):
        """Risks JSON output is valid and complete."""
        assert main(["ingest", str(SAMPLE_DIR)]) == 0
        capsys.readouterr()
        assert main(["risks", "--json"]) == 0
        output = capsys.readouterr().out
        data = json.loads(output)
        assert data["portfolio_rag"] in ("Red", "Amber", "Green")
        assert len(data["project_summaries"]) == 6

    def test_multiple_scenarios_in_sequence(self, tmp_path, capsys):
        """Multiple scenarios can run in sequence."""
        assert main(["ingest", str(SAMPLE_DIR)]) == 0

        scenarios = [
            "cut Beta scope by 30%",
            "increase Gamma budget by 50%",
            "delay Alpha by 1 quarter",
            "remove Delta",
        ]
        for text in scenarios:
            assert main(["scenario", text]) == 0

    def test_cross_format_consistency(self, tmp_path, capsys):
        """CSV and JSON ingestion produce consistent project counts."""
        # Ingest CSV
        csv_file = SAMPLE_DIR / "jira-export-sample.csv"
        csv_dir = tmp_path / "csv_only"
        csv_dir.mkdir()
        import shutil
        shutil.copy(csv_file, csv_dir / "data.csv")

        assert main(["ingest", str(csv_dir)]) == 0
        csv_projects = len(_session.projects)
        csv_risks = _session.report.total_risks

        # Reset and ingest JSON
        _session.projects = []
        json_file = SAMPLE_DIR / "jira-export-sample.json"
        json_dir = tmp_path / "json_only"
        json_dir.mkdir()
        shutil.copy(json_file, json_dir / "data.json")

        assert main(["ingest", str(json_dir)]) == 0
        json_projects = len(_session.projects)
        json_risks = _session.report.total_risks

        assert csv_projects == json_projects
        assert csv_risks == json_risks
