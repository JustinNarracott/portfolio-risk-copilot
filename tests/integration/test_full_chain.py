"""
Full-chain E2E tests — Issue #36.

Tests the complete pipeline: ingest (projects + benefits) → risks →
scenario → decisions → all artefacts.
"""

from datetime import date
from pathlib import Path

import pytest
from docx import Document

from src.cli import main, _session, Session


SAMPLE_DIR = Path(__file__).parent.parent.parent / "sample-data"
REF_DATE = date(2026, 2, 19)


@pytest.fixture(autouse=True)
def reset_session():
    """Reset session before each test."""
    _session.__init__()
    _session.reference_date = REF_DATE
    yield
    _session.__init__()


class TestFullPipeline:

    def test_ingest_loads_projects_and_benefits(self):
        rc = main(["ingest", str(SAMPLE_DIR)])
        assert rc == 0
        assert len(_session.projects) >= 11  # portfolio-export.csv has 11
        assert len(_session.benefits) >= 6
        assert _session.report is not None
        assert _session.benefit_report is not None
        assert _session.investment_report is not None
        assert len(_session.decision_log.decisions) >= 1

    def test_risks_after_ingest(self):
        main(["ingest", str(SAMPLE_DIR)])
        rc = main(["risks"])
        assert rc == 0

    def test_risks_json_after_ingest(self):
        main(["ingest", str(SAMPLE_DIR)])
        rc = main(["risks", "--json"])
        assert rc == 0

    def test_scenario_after_ingest(self):
        main(["ingest", str(SAMPLE_DIR)])
        rc = main(["scenario", "delay Alpha by 3 months"])
        assert rc == 0
        # Should have added a scenario decision
        assert any(d.source == "scenario" for d in _session.decision_log.decisions)

    def test_brief_all_generates_7_artefacts(self, tmp_path):
        main(["ingest", str(SAMPLE_DIR)])
        rc = main(["brief", "all", "--output-dir", str(tmp_path)])
        assert rc == 0
        docx_files = list(tmp_path.glob("*.docx"))
        pptx_files = list(tmp_path.glob("*.pptx"))
        assert len(docx_files) >= 5  # board, steering, project, benefits, investment, decisions
        assert len(pptx_files) >= 1  # board slides

    def test_board_briefing_has_action_summary(self, tmp_path):
        main(["ingest", str(SAMPLE_DIR)])
        main(["brief", "board", "--output-dir", str(tmp_path)])
        doc = Document(str(tmp_path / "board-briefing.docx"))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + "\n"
        assert "ACTION REQUIRED" in all_text or "PORTFOLIO SUMMARY" in all_text or "urgent" in all_text.lower()

    def test_steering_pack_has_benefits_section(self, tmp_path):
        main(["ingest", str(SAMPLE_DIR)])
        main(["brief", "steering", "--output-dir", str(tmp_path)])
        doc = Document(str(tmp_path / "steering-committee-pack.docx"))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # Check for table cell text too
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + "\n"
        assert "benefit" in all_text.lower() or "drift" in all_text.lower()

    def test_benefits_report_has_projects(self, tmp_path):
        main(["ingest", str(SAMPLE_DIR)])
        main(["brief", "benefits", "--output-dir", str(tmp_path)])
        doc = Document(str(tmp_path / "benefits-report.docx"))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + "\n"
        assert "£" in all_text

    def test_investment_summary_has_roi(self, tmp_path):
        main(["ingest", str(SAMPLE_DIR)])
        main(["brief", "investment", "--output-dir", str(tmp_path)])
        doc = Document(str(tmp_path / "investment-summary.docx"))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + "\n"
        assert "ROI" in all_text or "Invest" in all_text or "Hold" in all_text

    def test_decision_log_has_entries(self, tmp_path):
        main(["ingest", str(SAMPLE_DIR)])
        main(["scenario", "cut Beta budget by 30%"])
        main(["brief", "decisions", "--output-dir", str(tmp_path)])
        doc = Document(str(tmp_path / "decision-log.docx"))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "DEC-" in all_text

    def test_full_chain_under_30_seconds(self, tmp_path):
        """Full pipeline should complete in <30 seconds."""
        import time
        start = time.time()
        main(["ingest", str(SAMPLE_DIR)])
        main(["risks"])
        main(["scenario", "delay Alpha by 2 months"])
        main(["brief", "all", "--output-dir", str(tmp_path)])
        elapsed = time.time() - start
        assert elapsed < 30, f"Full chain took {elapsed:.1f}s (limit: 30s)"


class TestEdgeCases:

    def test_brief_without_ingest_fails(self):
        rc = main(["brief", "board"])
        assert rc == 1

    def test_risks_without_ingest_fails(self):
        rc = main(["risks"])
        assert rc == 1

    def test_scenario_without_ingest_fails(self):
        rc = main(["scenario", "delay Alpha by 1 month"])
        assert rc == 1

    def test_benefits_brief_without_benefit_data(self, tmp_path):
        """If only project data (no benefit tracker), benefits brief should fail gracefully."""
        # Ingest only the jira export (no benefit tracker)
        from src.ingestion.parser import parse_file
        projects = parse_file(SAMPLE_DIR / "jira-export-sample.csv")
        _session.projects = projects
        from src.risk_engine.engine import analyse_portfolio
        _session.report = analyse_portfolio(projects, top_n=5, reference_date=REF_DATE)
        from src.scenario.graph import build_dependency_graph
        _session.graph = build_dependency_graph(projects)
        # No benefit report
        _session.benefit_report = None
        rc = main(["brief", "benefits", "--output-dir", str(tmp_path)])
        assert rc == 1
