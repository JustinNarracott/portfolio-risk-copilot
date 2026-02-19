"""
Benefits artefact generator — standalone benefits report and steering pack integration.

Sprint 5 — Issue #31.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from src.artefacts.docx_generator import (
    BrandConfig,
    _apply_base_styles, _add_header_bar, _add_section_heading, _add_footer,
    _maybe_add_logo, _set_cell_bg, _set_cell_margins, _remove_table_borders,
    _set_table_borders, _highlight_run, _add_decision_item, _h,
    RAG_COLOURS, RAG_BG, RAG_DARK,
)
from src.benefits.calculator import (
    PortfolioBenefitReport, ProjectBenefitSummary, DriftRAG,
)

DRIFT_BG = {"Red": "F5B7B1", "Amber": "FAD7A0", "Green": "A9DFBF"}
DRIFT_TEXT = {"Red": "922B21", "Amber": "935116", "Green": "1E8449"}


def generate_benefits_report(
    benefit_report: PortfolioBenefitReport,
    brand: BrandConfig | None = None,
    output_path: str | Path | None = None,
) -> Path:
    """Generate standalone benefits realisation report (1-2 pages)."""
    brand = brand or BrandConfig()
    doc = Document()
    _apply_base_styles(doc, brand)
    _add_header_bar(doc, brand, _h(brand, "benefits_title", "Benefits Realisation Report"))
    _maybe_add_logo(doc, brand)

    # Benefits charts
    try:
        from src.charts import chart_benefits_waterfall, chart_benefits_drift
        from docx.shared import Inches as _Inches
        wf = chart_benefits_waterfall(benefit_report)
        doc.add_picture(str(wf), width=_Inches(5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    except Exception:
        pass

    # Benefits dashboard
    _add_benefits_dashboard(doc, benefit_report, brand)

    # Benefits RAG table
    _add_section_heading(doc, brand, "Benefits by Project")
    _add_benefits_table(doc, benefit_report, brand)

    # Top benefits at risk
    if benefit_report.top_benefits_at_risk:
        _add_section_heading(doc, brand, "Benefits at Risk")
        for b in benefit_report.top_benefits_at_risk[:5]:
            if b.expected_value > 0 or b.notes:
                _add_benefit_risk_card(doc, b, brand)

    # Drift analysis
    drifting = [s for s in benefit_report.project_summaries if s.drift_pct > 0.15]
    if drifting:
        _add_section_heading(doc, brand, "Benefits Drift Analysis")
        # Drift chart
        try:
            from src.charts import chart_benefits_drift
            drift_chart = chart_benefits_drift(benefit_report)
            doc.add_picture(str(drift_chart), width=Inches(5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        except Exception:
            pass
        for s in sorted(drifting, key=lambda x: -x.drift_pct):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            # Drift badge
            badge_bg = DRIFT_BG.get(s.drift_rag, "F0F0F0")
            badge_text_col = DRIFT_TEXT.get(s.drift_rag, "333333")
            badge = p.add_run(f" {s.drift_rag} ")
            badge.font.size = Pt(8)
            badge.font.bold = True
            badge.font.color.rgb = RGBColor.from_string(badge_text_col)
            _highlight_run(badge, badge_bg)
            name = p.add_run(f"  {s.project_name}: ")
            name.font.bold = True
            name.font.size = Pt(10)
            exp = p.add_run(s.drift_explanation)
            exp.font.size = Pt(9)
            exp.font.color.rgb = RGBColor(0x50, 0x50, 0x50)

    # Recommendations
    _add_section_heading(doc, brand, "Recommendations")
    for i, rec in enumerate(benefit_report.recommendations, 1):
        _add_decision_item(doc, rec, i, brand)

    _add_footer(doc)
    path = Path(output_path) if output_path else Path("benefits-report.docx")
    doc.save(str(path))
    return path


def add_benefits_to_steering(doc: Document, benefit_report: PortfolioBenefitReport, brand: BrandConfig) -> None:
    """Add a benefits section to an existing steering pack document."""
    _add_section_heading(doc, brand, _h(brand, "benefits_section", "Benefits & Value Realisation"))

    # Compact dashboard
    _add_benefits_dashboard(doc, benefit_report, brand, compact=True)

    # Brief drift summary
    drifting = [s for s in benefit_report.project_summaries if s.drift_pct > 0.15]
    if drifting:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        b = p.add_run(f"{len(drifting)} project{'s' if len(drifting) > 1 else ''} showing benefits drift: ")
        b.font.bold = True
        b.font.size = Pt(10)
        names = ", ".join(f"{s.project_name} ({s.drift_pct:.0%})" for s in sorted(drifting, key=lambda x: -x.drift_pct)[:4])
        n = p.add_run(names)
        n.font.size = Pt(10)
        n.font.color.rgb = RGBColor(0x50, 0x50, 0x50)

    # Top recommendation
    if benefit_report.recommendations:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        arrow = p.add_run("→ ")
        arrow.font.bold = True
        arrow.font.color.rgb = RGBColor.from_string(brand.accent_colour)
        arrow.font.size = Pt(9)
        r = p.add_run(benefit_report.recommendations[0])
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor.from_string(brand.accent_colour)


# ──────────────────────────────────────────────
# Components
# ──────────────────────────────────────────────

def _add_benefits_dashboard(doc: Document, report: PortfolioBenefitReport, brand: BrandConfig, compact: bool = False) -> None:
    """Visual dashboard with benefits KPIs."""
    cols = 4
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(table)

    drift_bg = DRIFT_BG.get(report.portfolio_drift_rag, brand.primary_colour)
    drift_text = DRIFT_TEXT.get(report.portfolio_drift_rag, "FFFFFF")

    stats = [
        (f"£{report.total_expected:,.0f}", "EXPECTED\nVALUE", brand.primary_colour, "FFFFFF"),
        (f"£{report.total_realised:,.0f}", "REALISED\nTO DATE", "27AE60" if report.total_realised > 0 else "BDC3C7", "FFFFFF"),
        (f"{report.realisation_pct:.0%}", "REALISATION\nRATE", brand.accent_colour, "FFFFFF"),
        (f"{report.portfolio_drift_pct:.0%}", "BENEFITS\nDRIFT", drift_bg, drift_text),
    ]

    size_val = Pt(18) if compact else Pt(22)
    size_label = Pt(6) if compact else Pt(7)

    for i, (value, label, bg, text_col) in enumerate(stats):
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, bg)
        _set_cell_margins(cell, 100, 100, 120, 120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value)
        run.font.size = size_val
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(text_col)
        run.font.name = brand.heading_font
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.font.size = size_label
        r2.font.bold = True
        r2.font.color.rgb = RGBColor.from_string(text_col)

    doc.add_paragraph()


def _add_benefits_table(doc: Document, report: PortfolioBenefitReport, brand: BrandConfig) -> None:
    """Benefits per-project table with drift RAG."""
    headers = ["Project", "Expected", "Realised", "Drift", "Status"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, brand.primary_colour)
        _set_cell_margins(cell, 50, 50, 80, 80)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for idx, s in enumerate(report.project_summaries):
        row = table.add_row()
        bg = "F8F9FA" if idx % 2 == 0 else "FFFFFF"

        row.cells[0].text = s.project_name
        row.cells[1].text = f"£{s.total_expected:,.0f}"
        row.cells[2].text = f"£{s.total_realised:,.0f}"

        # Drift cell with colour
        drift_cell = row.cells[3]
        drift_cell.text = ""
        p = drift_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        drift_text = f"{s.drift_pct:.0%}" if s.total_expected > 0 else "N/A"
        run = p.add_run(f" {drift_text} ")
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        d_text_col = DRIFT_TEXT.get(s.drift_rag, "333333")
        d_bg = DRIFT_BG.get(s.drift_rag, "F0F0F0")
        run.font.color.rgb = RGBColor.from_string(d_text_col)
        _set_cell_bg(drift_cell, d_bg)

        # Status — count of benefits and their states
        ben_statuses = [b.status.value for b in s.benefits]
        row.cells[4].text = ", ".join(ben_statuses) if ben_statuses else "—"

        for cell in row.cells:
            if cell != drift_cell:
                _set_cell_bg(cell, bg)
            _set_cell_margins(cell, 40, 40, 80, 80)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.color.rgb is None:
                        run.font.size = Pt(9)

    _set_table_borders(table, "D5D8DC")


def _add_benefit_risk_card(doc: Document, b, brand: BrandConfig) -> None:
    """Single benefit at risk as a card."""
    from src.benefits.parser import Benefit
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)

    # Value badge
    if b.expected_value > 0:
        val_run = p.add_run(f" £{b.unrealised_value:,.0f} at risk ")
        val_run.font.size = Pt(8)
        val_run.font.bold = True
        val_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _highlight_run(val_run, "C0392B")

    name_run = p.add_run(f"  {b.name}")
    name_run.font.bold = True
    name_run.font.size = Pt(10)

    proj_run = p.add_run(f"  [{b.project_name}]")
    proj_run.font.size = Pt(9)
    proj_run.font.italic = True
    proj_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    if b.notes:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.3)
        p2.paragraph_format.space_after = Pt(6)
        r = p2.add_run(b.notes)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
