"""
Portfolio investment summary artefacts — DOCX report and PPTX slide.

Sprint 6 — Issue #34.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from src.artefacts.docx_generator import (
    BrandConfig,
    _apply_base_styles, _add_header_bar, _add_section_heading, _add_footer,
    _maybe_add_logo, _set_cell_bg, _set_cell_margins, _remove_table_borders,
    _set_table_borders, _highlight_run, _add_decision_item, _h,
)
from src.investment import (
    PortfolioInvestmentReport, InvestmentAction,
)

ACTION_COLOURS = {
    InvestmentAction.INVEST: ("1E8449", "D5F5E3"),
    InvestmentAction.HOLD: ("2471A3", "D6EAF8"),
    InvestmentAction.REVIEW: ("B7950B", "FEF9E7"),
    InvestmentAction.DIVEST: ("922B21", "FADBD8"),
}


def generate_investment_report(
    investment_report: PortfolioInvestmentReport,
    brand: BrandConfig | None = None,
    output_path: str | Path | None = None,
) -> Path:
    """Generate standalone investment summary DOCX."""
    brand = brand or BrandConfig()
    doc = Document()
    _apply_base_styles(doc, brand)
    _add_header_bar(doc, brand, _h(brand, "investment_title", "Portfolio Investment Summary"))
    _maybe_add_logo(doc, brand)

    # Investment dashboard
    _add_investment_dashboard(doc, investment_report, brand)

    # ROI league table
    _add_section_heading(doc, brand, "Investment Ranking (by ROI)")
    _add_roi_table(doc, investment_report, brand)

    # Invest/Hold/Divest summary
    _add_section_heading(doc, brand, "Invest / Hold / Divest")
    _add_action_summary(doc, investment_report, brand)

    # Value at risk
    if investment_report.top_value_at_risk:
        _add_section_heading(doc, brand, "Value at Risk")
        for pi in investment_report.top_value_at_risk:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            badge_text, badge_bg = ACTION_COLOURS.get(pi.action, ("333333", "F0F0F0"))
            badge = p.add_run(f" {pi.action.value} ")
            badge.font.size = Pt(8)
            badge.font.bold = True
            badge.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _highlight_run(badge, badge_text)
            name = p.add_run(f"  {pi.project_name}")
            name.font.bold = True
            name.font.size = Pt(10)
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.3)
            p2.paragraph_format.space_after = Pt(6)
            r = p2.add_run(pi.action_rationale)
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x50, 0x50, 0x50)

    # Recommendations
    _add_section_heading(doc, brand, "Recommendations")
    for i, rec in enumerate(investment_report.recommendations, 1):
        _add_decision_item(doc, rec, i, brand)

    _add_footer(doc)
    path = Path(output_path) if output_path else Path("investment-summary.docx")
    doc.save(str(path))
    return path


# ──────────────────────────────────────────────
# Components
# ──────────────────────────────────────────────

def _add_investment_dashboard(doc: Document, report: PortfolioInvestmentReport, brand: BrandConfig) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(table)

    roi_bg = "27AE60" if report.portfolio_roi > 0 else "C0392B"

    stats = [
        (f"£{report.total_budget:,.0f}", "TOTAL\nBUDGET", brand.primary_colour),
        (f"£{report.total_spent:,.0f}", "TOTAL\nSPENT", brand.accent_colour),
        (f"£{report.total_cost_to_complete:,.0f}", "COST TO\nCOMPLETE", "7F8C8D"),
        (f"{report.portfolio_roi:.0%}", "PORTFOLIO\nROI", roi_bg),
    ]

    for i, (value, label, bg) in enumerate(stats):
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, bg)
        _set_cell_margins(cell, 100, 100, 120, 120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value)
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = brand.heading_font
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.font.size = Pt(7)
        r2.font.bold = True
        r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()


def _add_roi_table(doc: Document, report: PortfolioInvestmentReport, brand: BrandConfig) -> None:
    headers = ["#", "Project", "Budget", "ROI", "Action", "RAG"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, brand.primary_colour)
        _set_cell_margins(cell, 50, 50, 80, 80)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for idx, pi in enumerate(report.project_investments):
        row = table.add_row()
        bg = "F8F9FA" if idx % 2 == 0 else "FFFFFF"
        row.cells[0].text = str(pi.roi_rank)
        row.cells[1].text = pi.project_name
        row.cells[2].text = f"£{pi.budget:,.0f}"
        row.cells[3].text = f"{pi.roi:.0%}"

        # Action cell with colour
        action_cell = row.cells[4]
        action_cell.text = ""
        p = action_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text_col, bg_col = ACTION_COLOURS.get(pi.action, ("333333", "F0F0F0"))
        run = p.add_run(f" {pi.action.value} ")
        run.font.bold = True
        run.font.size = Pt(8)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(text_col)
        _set_cell_bg(action_cell, bg_col)

        # RAG cell
        from src.artefacts.docx_generator import RAG_BG, RAG_COLOURS
        rag_cell = row.cells[5]
        rag_cell.text = ""
        p = rag_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f" {pi.rag_status} ")
        run.font.bold = True
        run.font.size = Pt(8)
        run.font.name = "Calibri"
        run.font.color.rgb = RAG_COLOURS.get(pi.rag_status, RGBColor(0x33, 0x33, 0x33))
        _set_cell_bg(rag_cell, RAG_BG.get(pi.rag_status, "F0F0F0"))

        for cell in row.cells:
            if cell not in (action_cell, rag_cell):
                _set_cell_bg(cell, bg)
            _set_cell_margins(cell, 40, 40, 80, 80)
            for paragraph in cell.paragraphs:
                for r in paragraph.runs:
                    if r.font.color.rgb is None:
                        r.font.size = Pt(9)

    _set_table_borders(table, "D5D8DC")


def _add_action_summary(doc: Document, report: PortfolioInvestmentReport, brand: BrandConfig) -> None:
    """Visual summary of Invest/Hold/Divest breakdown."""
    groups = {}
    for pi in report.project_investments:
        groups.setdefault(pi.action, []).append(pi)

    for action in [InvestmentAction.INVEST, InvestmentAction.HOLD, InvestmentAction.REVIEW, InvestmentAction.DIVEST]:
        items = groups.get(action, [])
        if not items:
            continue
        text_col, bg_col = ACTION_COLOURS.get(action, ("333333", "F0F0F0"))
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        badge = p.add_run(f" {action.value.upper()} ({len(items)}) ")
        badge.font.size = Pt(9)
        badge.font.bold = True
        badge.font.color.rgb = RGBColor.from_string(text_col)
        _highlight_run(badge, bg_col)
        names = p.add_run(f"  {', '.join(pi.project_name for pi in items)}")
        names.font.size = Pt(10)

        # Total budget in this group
        total = sum(pi.budget for pi in items)
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.3)
        p2.paragraph_format.space_after = Pt(6)
        r = p2.add_run(f"Combined budget: £{total:,.0f}")
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
