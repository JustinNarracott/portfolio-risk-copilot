"""
Decision log generator.

Captures decisions from scenario simulations, risk analysis, and
investment recommendations into a structured audit trail.

Sprint 6 — Issue #33.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import date, datetime
from enum import Enum
from pathlib import Path
from typing import Any
import json

from src.scenario.simulator import ScenarioResult, ProjectImpact
from src.scenario.narrative import generate_narrative, ScenarioNarrative
from src.risk_engine.engine import PortfolioRiskReport
from src.investment import PortfolioInvestmentReport, InvestmentAction


class DecisionStatus(Enum):
    PENDING = "Pending"
    APPROVED = "Approved"
    REJECTED = "Rejected"
    DEFERRED = "Deferred"


@dataclass
class DecisionOption:
    """A single option considered for a decision."""
    label: str
    description: str
    impact_summary: str


@dataclass
class Decision:
    """A single decision record."""
    decision_id: str
    date: str
    title: str
    context: str
    projects_affected: list[str]
    options: list[DecisionOption]
    recommendation: str
    recommendation_rationale: str
    status: DecisionStatus
    source: str  # "scenario", "risk_analysis", "investment_review"

    def to_dict(self) -> dict[str, Any]:
        return {
            "decision_id": self.decision_id,
            "date": self.date,
            "title": self.title,
            "context": self.context,
            "projects_affected": self.projects_affected,
            "options": [{"label": o.label, "description": o.description, "impact": o.impact_summary} for o in self.options],
            "recommendation": self.recommendation,
            "recommendation_rationale": self.recommendation_rationale,
            "status": self.status.value,
            "source": self.source,
        }


@dataclass
class DecisionLog:
    """Complete decision log for a portfolio review cycle."""
    decisions: list[Decision] = field(default_factory=list)
    _counter: int = field(default=0, repr=False)

    def add(self, decision: Decision) -> None:
        self.decisions.append(decision)

    def next_id(self) -> str:
        self._counter += 1
        return f"DEC-{self._counter:03d}"

    def to_dict(self) -> dict[str, Any]:
        return {
            "decision_count": len(self.decisions),
            "decisions": [d.to_dict() for d in self.decisions],
        }

    def to_json(self, indent: int = 2) -> str:
        return json.dumps(self.to_dict(), indent=indent)


# ──────────────────────────────────────────────
# Decision generators
# ──────────────────────────────────────────────

def decision_from_scenario(
    result: ScenarioResult, log: DecisionLog, ref_date: date | None = None,
) -> Decision:
    """Create a decision record from a scenario simulation result."""
    ref = ref_date or date.today()
    narrative = generate_narrative(result)
    action_desc = result.action.description or f"{result.action.action.value} on {result.action.project}"

    # Collect affected projects
    projects = list({i.project_name for i in result.impacts})
    if not projects:
        projects = [result.action.project]

    # Build options: the scenario action vs do-nothing
    option_a = DecisionOption(
        label=f"Apply: {action_desc}",
        description=narrative.impact_analysis or narrative.after_summary,
        impact_summary=narrative.after_summary,
    )
    option_b = DecisionOption(
        label="Do nothing — maintain current plan",
        description="Continue on current trajectory without change.",
        impact_summary=narrative.before_summary or "No change to delivery dates, budget, or dependencies.",
    )

    # Recommendation based on whether scenario has warnings
    if result.warnings:
        rec = f"Proceed with caution — {action_desc}"
        rationale = f"Scenario modelled successfully but {len(result.warnings)} warning(s) flagged: {'; '.join(result.warnings[:2])}."
    else:
        rec = f"Recommend: {action_desc}"
        rationale = "; ".join(narrative.recommendations) if narrative.recommendations else "Scenario impact is manageable."

    decision = Decision(
        decision_id=log.next_id(),
        date=ref.isoformat(),
        title=f"Scenario: {action_desc}",
        context=f"Scenario simulation for {result.action.project}.",
        projects_affected=projects,
        options=[option_a, option_b],
        recommendation=rec,
        recommendation_rationale=rationale,
        status=DecisionStatus.PENDING,
        source="scenario",
    )
    log.add(decision)
    return decision


def decisions_from_risk_report(
    report: PortfolioRiskReport, log: DecisionLog, ref_date: date | None = None,
) -> list[Decision]:
    """Generate decisions from critical risk findings."""
    ref = ref_date or date.today()
    decisions: list[Decision] = []

    red_projects = [s for s in report.project_summaries if s.rag_status == "Red"]
    if red_projects:
        names = [s.project_name for s in red_projects]
        d = Decision(
            decision_id=log.next_id(),
            date=ref.isoformat(),
            title=f"Escalate {len(red_projects)} Red project{'s' if len(red_projects) > 1 else ''} to executive review",
            context=f"{len(red_projects)} projects at Red status with combined {sum(s.risk_count for s in red_projects)} risks.",
            projects_affected=names[:5],
            options=[
                DecisionOption("Escalate to executive review", "Schedule emergency review within 5 days.", "Leadership intervention, possible resource reallocation."),
                DecisionOption("Enhanced monitoring", "Increase reporting frequency to weekly.", "Earlier detection but no direct intervention."),
                DecisionOption("Accept risk", "Continue with current oversight level.", "No additional overhead but risk of further deterioration."),
            ],
            recommendation="Escalate to executive review",
            recommendation_rationale=f"{len(red_projects)} projects at Red status requires leadership attention — monitoring alone is insufficient.",
            status=DecisionStatus.PENDING,
            source="risk_analysis",
        )
        log.add(d)
        decisions.append(d)

    return decisions


def decisions_from_investment(
    investment_report: PortfolioInvestmentReport, log: DecisionLog, ref_date: date | None = None,
) -> list[Decision]:
    """Generate decisions from investment analysis."""
    ref = ref_date or date.today()
    decisions: list[Decision] = []

    divests = [p for p in investment_report.project_investments if p.action == InvestmentAction.DIVEST]
    if divests:
        freed = sum(p.cost_to_complete for p in divests)
        names = [p.project_name for p in divests]
        invests = [p for p in investment_report.project_investments if p.action == InvestmentAction.INVEST]
        invest_names = ", ".join(p.project_name for p in invests[:3]) if invests else "higher-value initiatives"

        d = Decision(
            decision_id=log.next_id(),
            date=ref.isoformat(),
            title=f"Divest from {', '.join(names[:2])} — reallocate £{freed:,.0f}",
            context=f"{len(divests)} project{'s' if len(divests) > 1 else ''} showing negative ROI with Red delivery status.",
            projects_affected=names,
            options=[
                DecisionOption("Full divestment", f"Stop discretionary spend on {', '.join(names[:2])}.", f"Frees £{freed:,.0f} for reallocation to {invest_names}."),
                DecisionOption("Reduced scope", "Cut scope to minimum viable and reduce budget.", "Partial savings, some benefit preserved."),
                DecisionOption("Continue as-is", "Maintain current investment level.", "No freed budget. Risk of further value erosion."),
            ],
            recommendation="Full divestment — redirect budget to higher-ROI projects",
            recommendation_rationale=f"Continuing to invest in negative-ROI, Red-status projects erodes portfolio value. £{freed:,.0f} better deployed on {invest_names}.",
            status=DecisionStatus.PENDING,
            source="investment_review",
        )
        log.add(d)
        decisions.append(d)

    return decisions


# ──────────────────────────────────────────────
# DOCX export
# ──────────────────────────────────────────────

def export_decision_log(
    log: DecisionLog, brand=None, output_path: str | Path | None = None,
) -> Path:
    """Export decision log as DOCX."""
    from src.artefacts.docx_generator import (
        BrandConfig, _apply_base_styles, _add_header_bar, _add_section_heading,
        _add_footer, _set_cell_bg, _set_cell_margins, _set_table_borders, _h,
    )
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT

    brand = brand or BrandConfig()
    doc = Document()
    _apply_base_styles(doc, brand)
    _add_header_bar(doc, brand, _h(brand, "decisions_title", "Portfolio Decision Log"))

    if not log.decisions:
        p = doc.add_paragraph()
        p.add_run("No decisions recorded this cycle.").font.italic = True
    else:
        for d in log.decisions:
            _add_section_heading(doc, brand, f"{d.decision_id}: {d.title}", level=2)

            # Context
            p = doc.add_paragraph()
            p.add_run("Context: ").font.bold = True
            p.add_run(d.context).font.size = Pt(10)

            p2 = doc.add_paragraph()
            p2.add_run("Projects: ").font.bold = True
            p2.add_run(", ".join(d.projects_affected)).font.size = Pt(10)

            p3 = doc.add_paragraph()
            p3.add_run("Source: ").font.bold = True
            p3.add_run(d.source.replace("_", " ").title()).font.size = Pt(10)

            # Options table
            headers = ["Option", "Description", "Impact"]
            table = doc.add_table(rows=1, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            for i, text in enumerate(headers):
                cell = table.rows[0].cells[i]
                _set_cell_bg(cell, brand.primary_colour)
                _set_cell_margins(cell, 40, 40, 80, 80)
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.font.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            for opt in d.options:
                row = table.add_row()
                row.cells[0].text = opt.label
                row.cells[1].text = opt.description
                row.cells[2].text = opt.impact_summary
                for cell in row.cells:
                    _set_cell_margins(cell, 30, 30, 60, 60)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
            _set_table_borders(table, "D5D8DC")

            # Recommendation
            p4 = doc.add_paragraph()
            p4.paragraph_format.space_before = Pt(6)
            rec = p4.add_run("→ Recommendation: ")
            rec.font.bold = True
            rec.font.size = Pt(10)
            rec.font.color.rgb = RGBColor.from_string(brand.accent_colour)
            p4.add_run(d.recommendation).font.size = Pt(10)

            # Status
            p5 = doc.add_paragraph()
            status_run = p5.add_run(f"Status: {d.status.value}")
            status_run.font.size = Pt(9)
            status_run.font.italic = True
            status_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    _add_footer(doc)
    path = Path(output_path) if output_path else Path("decision-log.docx")
    doc.save(str(path))
    return path
