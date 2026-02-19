"""
DOCX document generator — professional exec-ready briefings.

Generates polished Word documents with visual hierarchy, colour-coded RAG,
severity badges, branded headers, and data-driven recommendations.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.risk_engine.engine import (
    PortfolioRiskReport,
    ProjectRiskSummary,
    Risk,
    RiskCategory,
    RiskSeverity,
)


@dataclass
class BrandConfig:
    """User-customisable branding for generated documents."""
    logo_path: str | None = None
    primary_colour: str = "1B3A5C"
    accent_colour: str = "2E75B6"
    heading_font: str = "Calibri"
    body_font: str = "Calibri"
    company_name: str = ""
    custom_headings: dict[str, str] = field(default_factory=dict)


RAG_COLOURS = {"Red": RGBColor(0xC0, 0x39, 0x2B), "Amber": RGBColor(0xE6, 0x7E, 0x22), "Green": RGBColor(0x27, 0xAE, 0x60)}
RAG_BG = {"Red": "F5B7B1", "Amber": "FAD7A0", "Green": "A9DFBF"}
RAG_DARK = {"Red": "C0392B", "Amber": "E67E22", "Green": "27AE60"}
SEVERITY_COLOURS = {
    RiskSeverity.CRITICAL: RGBColor(0xC0, 0x39, 0x2B),
    RiskSeverity.HIGH: RGBColor(0xE6, 0x7E, 0x22),
    RiskSeverity.MEDIUM: RGBColor(0xF3, 0x9C, 0x12),
    RiskSeverity.LOW: RGBColor(0x27, 0xAE, 0x60),
}
SEVERITY_BADGES = {
    RiskSeverity.CRITICAL: ("CRITICAL", "C0392B"),
    RiskSeverity.HIGH: ("HIGH", "E67E22"),
    RiskSeverity.MEDIUM: ("MEDIUM", "F39C12"),
    RiskSeverity.LOW: ("LOW", "27AE60"),
}


# ──────────────────────────────────────────────
# Public API
# ──────────────────────────────────────────────

def generate_board_briefing(
    report: PortfolioRiskReport, brand: BrandConfig | None = None, output_path: str | Path | None = None,
    benefit_report=None, investment_report=None,
) -> Path:
    """Generate a 1-page board briefing DOCX."""
    brand = brand or BrandConfig()
    doc = Document()
    _apply_base_styles(doc, brand)
    _add_header_bar(doc, brand, _h(brand, "board_title", "Portfolio Health — Board Briefing"))
    _maybe_add_logo(doc, brand)
    # Executive action summary — the "7am phone check" paragraph
    from src.insights import generate_executive_summary
    exec_summary = generate_executive_summary(report, benefit_report, investment_report)
    _add_exec_action_box(doc, exec_summary, brand)
    _add_portfolio_dashboard(doc, report, brand)
    _add_section_heading(doc, brand, _h(brand, "project_table", "Project Overview"))
    _add_project_rag_table(doc, report, brand)
    _add_section_heading(doc, brand, _h(brand, "top_risks", "Top Risks"))
    for i, risk in enumerate(_get_top_n_risks(report, 3), 1):
        _add_risk_card(doc, risk, index=i)
    _add_section_heading(doc, brand, _h(brand, "decisions", "Recommended Decisions"))
    for i, d in enumerate(_generate_decisions(report, 3), 1):
        _add_decision_item(doc, d, i, brand)
    _add_footer(doc)
    path = Path(output_path) if output_path else Path("board-briefing.docx")
    doc.save(str(path))
    return path


def generate_steering_pack(
    report: PortfolioRiskReport, brand: BrandConfig | None = None, output_path: str | Path | None = None,
    benefit_report=None, investment_report=None,
) -> Path:
    """Generate a 2-3 page steering committee pack DOCX."""
    brand = brand or BrandConfig()
    doc = Document()
    _apply_base_styles(doc, brand)
    _add_header_bar(doc, brand, _h(brand, "steering_title", "Portfolio Risk & Value Briefing — Steering Committee"))
    _maybe_add_logo(doc, brand)
    # Executive action summary — the lead paragraph
    from src.insights import generate_executive_summary
    exec_summary = generate_executive_summary(report, benefit_report, investment_report)
    _add_exec_action_box(doc, exec_summary, brand)
    _add_section_heading(doc, brand, _h(brand, "exec_summary", "Executive Summary"))
    _add_exec_summary(doc, report, brand)
    _add_portfolio_dashboard(doc, report, brand)
    _add_section_heading(doc, brand, _h(brand, "project_table", "Project Overview"))
    _add_project_rag_table(doc, report, brand, detailed=True)
    _add_section_heading(doc, brand, _h(brand, "top_risks", "Top Portfolio Risks"))
    for i, risk in enumerate(_get_top_n_risks(report, 5), 1):
        _add_risk_card(doc, risk, index=i, include_mitigation=True)
    _add_section_heading(doc, brand, _h(brand, "decisions", "Recommended Decisions"))
    for i, d in enumerate(_generate_decisions(report, 3), 1):
        _add_decision_item(doc, d, i, brand)
    # Benefits section (if benefit data available)
    if benefit_report is not None:
        from src.benefits.artefacts import add_benefits_to_steering
        add_benefits_to_steering(doc, benefit_report, brand)
    _add_section_heading(doc, brand, _h(brand, "talking_points", "Talking Points for Discussion"))
    for pt in _generate_talking_points(report):
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(pt)
        r.font.size = Pt(10)
        r.font.name = brand.body_font
    _add_section_heading(doc, brand, "Risk Distribution by Category")
    _add_risk_distribution_table(doc, report, brand)
    _add_footer(doc)
    path = Path(output_path) if output_path else Path("steering-committee-pack.docx")
    doc.save(str(path))
    return path


def generate_project_status_pack(
    report: PortfolioRiskReport, brand: BrandConfig | None = None, output_path: str | Path | None = None,
) -> Path:
    """Generate per-project status pack DOCX."""
    brand = brand or BrandConfig()
    doc = Document()
    _apply_base_styles(doc, brand)
    _add_header_bar(doc, brand, _h(brand, "status_title", "Project Status Pack"))
    _maybe_add_logo(doc, brand)
    for idx, summary in enumerate(report.project_summaries):
        if idx > 0:
            doc.add_page_break()
        _add_project_header(doc, summary, brand)
        if summary.risks:
            _add_section_heading(doc, brand, "Risks", level=2)
            for risk in summary.risks:
                _add_risk_card(doc, risk, include_mitigation=True)
        else:
            p = doc.add_paragraph()
            p.add_run("No significant risks identified.").font.italic = True
        _add_section_heading(doc, brand, "Immediate Actions", level=2)
        actions = _generate_project_actions(summary)
        if actions:
            for i, a in enumerate(actions, 1):
                _add_decision_item(doc, a, i, brand)
        else:
            p = doc.add_paragraph()
            r = p.add_run("Continue on current trajectory. No escalation needed.")
            r.font.italic = True
            r.font.size = Pt(10)
    _add_footer(doc)
    path = Path(output_path) if output_path else Path("project-status-pack.docx")
    doc.save(str(path))
    return path


# ──────────────────────────────────────────────
# Layout components
# ──────────────────────────────────────────────

def _add_exec_action_box(doc: Document, text: str, brand: BrandConfig) -> None:
    """Visually distinct action summary callout box."""
    # Determine box colour based on content
    has_urgent = "urgent" in text.lower() or "emergency" in text.lower()
    bg = "FEF5E7" if has_urgent else "EBF5FB"
    border_col = "E67E22" if has_urgent else brand.accent_colour

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    _set_cell_bg(cell, bg)
    _set_cell_margins(cell, 150, 150, 200, 200)

    # Label
    p = cell.paragraphs[0]
    label = p.add_run("⚡ ACTION REQUIRED  " if has_urgent else "📋 PORTFOLIO SUMMARY  ")
    label.font.size = Pt(8)
    label.font.bold = True
    label.font.color.rgb = RGBColor.from_string(border_col)

    # Content
    p2 = cell.add_paragraph()
    run = p2.add_run(text)
    run.font.size = Pt(10)
    run.font.name = brand.body_font
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    _set_table_borders(table, border_col)
    doc.add_paragraph()


def _apply_base_styles(doc: Document, brand: BrandConfig) -> None:
    style = doc.styles["Normal"]
    style.font.name = brand.body_font
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)


def _add_header_bar(doc: Document, brand: BrandConfig, text: str) -> None:
    """Coloured header bar with white text — simulated via single-cell table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    _set_cell_bg(cell, brand.primary_colour)
    _set_cell_margins(cell, 150, 150, 200, 200)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = brand.heading_font
    _remove_table_borders(table)
    doc.add_paragraph()


def _add_section_heading(doc: Document, brand: BrandConfig, text: str, level: int = 1) -> None:
    """Section heading with accent underline."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    # Bottom border
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), brand.accent_colour)
    bottom.set(qn("w:space"), "1")
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.size = Pt(14) if level == 1 else Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(brand.primary_colour)
    run.font.name = brand.heading_font


def _add_portfolio_dashboard(doc: Document, report: PortfolioRiskReport, brand: BrandConfig) -> None:
    """Visual dashboard — key stats in coloured boxes."""
    total = len(report.project_summaries)
    reds = sum(1 for s in report.project_summaries if s.rag_status == "Red")
    ambers = sum(1 for s in report.project_summaries if s.rag_status == "Amber")

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(table)

    stats = [
        (report.portfolio_rag, "PORTFOLIO\nSTATUS", RAG_DARK.get(report.portfolio_rag, brand.primary_colour)),
        (str(reds), "RED\nPROJECTS", "C0392B" if reds > 0 else "BDC3C7"),
        (str(ambers), "AMBER\nPROJECTS", "E67E22" if ambers > 0 else "BDC3C7"),
        (str(report.total_risks), "TOTAL\nRISKS", brand.primary_colour),
    ]
    for i, (value, label, bg) in enumerate(stats):
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, bg)
        _set_cell_margins(cell, 120, 120, 150, 150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value)
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = brand.heading_font
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.font.size = Pt(7)
        r2.font.bold = True
        r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()


def _add_project_rag_table(doc: Document, report: PortfolioRiskReport, brand: BrandConfig, detailed: bool = False) -> None:
    """Polished RAG table with coloured status cells."""
    headers = ["Project", "Status", "RAG", "Risks", "Top Risk"] if detailed else ["Project", "Status", "RAG", "Risks"]
    col_count = len(headers)
    table = doc.add_table(rows=1, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, brand.primary_colour)
        _set_cell_margins(cell, 60, 60, 100, 100)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for idx, s in enumerate(report.project_summaries):
        row = table.add_row()
        bg = "F8F9FA" if idx % 2 == 0 else "FFFFFF"
        row.cells[0].text = s.project_name
        row.cells[1].text = s.project_status
        # RAG cell — dark text on light coloured background
        rag_cell = row.cells[2]
        rag_cell.text = ""
        p = rag_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"  {s.rag_status}  ")
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        rag_text_col = RAG_COLOURS.get(s.rag_status, RGBColor(0x33, 0x33, 0x33))
        run.font.color.rgb = rag_text_col
        _set_cell_bg(rag_cell, RAG_BG.get(s.rag_status, "F0F0F0"))
        row.cells[3].text = str(s.risk_count)
        if detailed and s.risks:
            row.cells[4].text = s.risks[0].title
        for cell in row.cells:
            if cell != rag_cell:
                _set_cell_bg(cell, bg)
            _set_cell_margins(cell, 50, 50, 100, 100)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.color.rgb is None:
                        run.font.size = Pt(9)
    _set_table_borders(table, "D5D8DC")


def _add_risk_card(doc: Document, risk: Risk, index: int | None = None, include_mitigation: bool = False) -> None:
    """Risk as a visually distinct card with severity badge."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)

    badge_text, badge_colour = SEVERITY_BADGES.get(risk.severity, ("?", "BDC3C7"))
    badge_run = p.add_run(f" {badge_text} ")
    badge_run.font.size = Pt(8)
    badge_run.font.bold = True
    badge_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _highlight_run(badge_run, badge_colour)

    prefix = f"  {index}. " if index else "  "
    title_run = p.add_run(f"{prefix}{risk.title}")
    title_run.font.bold = True
    title_run.font.size = Pt(10)
    title_run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    proj_run = p.add_run(f"  [{risk.project_name}]")
    proj_run.font.size = Pt(9)
    proj_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    proj_run.font.italic = True

    # Explanation
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.3)
    p2.paragraph_format.space_after = Pt(2)
    r = p2.add_run(risk.explanation)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x50, 0x50, 0x50)

    if include_mitigation and risk.suggested_mitigation:
        p3 = doc.add_paragraph()
        p3.paragraph_format.left_indent = Inches(0.3)
        p3.paragraph_format.space_after = Pt(8)
        arrow = p3.add_run("→ ")
        arrow.font.size = Pt(9)
        arrow.font.bold = True
        arrow.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        mit = p3.add_run(risk.suggested_mitigation)
        mit.font.size = Pt(9)
        mit.font.italic = True
        mit.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)


def _add_decision_item(doc: Document, text: str, index: int, brand: BrandConfig) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    num = p.add_run(f" {index} ")
    num.font.size = Pt(9)
    num.font.bold = True
    num.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _highlight_run(num, brand.accent_colour)
    dec = p.add_run(f"  {text}")
    dec.font.size = Pt(10)
    dec.font.name = brand.body_font


def _add_project_header(doc: Document, summary: ProjectRiskSummary, brand: BrandConfig) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    name = p.add_run(summary.project_name)
    name.font.size = Pt(14)
    name.font.bold = True
    name.font.color.rgb = RGBColor.from_string(brand.primary_colour)
    name.font.name = brand.heading_font
    rag_run = p.add_run(f"  {summary.rag_status}  ")
    rag_run.font.size = Pt(9)
    rag_run.font.bold = True
    rag_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _highlight_run(rag_run, RAG_DARK.get(summary.rag_status, "BDC3C7"))
    st = p.add_run(f"  {summary.project_status}")
    st.font.size = Pt(10)
    st.font.color.rgb = RGBColor(0x70, 0x70, 0x70)


def _add_exec_summary(doc: Document, report: PortfolioRiskReport, brand: BrandConfig) -> None:
    total = len(report.project_summaries)
    reds = sum(1 for s in report.project_summaries if s.rag_status == "Red")
    ambers = sum(1 for s in report.project_summaries if s.rag_status == "Amber")
    greens = sum(1 for s in report.project_summaries if s.rag_status == "Green")
    p = doc.add_paragraph()
    p.add_run(
        f"The portfolio comprises {total} active projects. "
        f"{reds} rated Red (immediate attention), {ambers} Amber (emerging risks), "
        f"and {greens} Green (on track). {report.total_risks} risks identified."
    ).font.size = Pt(10)
    if reds > 0:
        red_names = [s.project_name for s in report.project_summaries if s.rag_status == "Red"]
        p2 = doc.add_paragraph()
        b = p2.add_run("Immediate attention: ")
        b.font.bold = True
        b.font.size = Pt(10)
        b.font.color.rgb = RAG_COLOURS["Red"]
        p2.add_run(", ".join(red_names)).font.size = Pt(10)


def _add_risk_distribution_table(doc: Document, report: PortfolioRiskReport, brand: BrandConfig) -> None:
    cats: dict[str, dict[str, int]] = {}
    for s in report.project_summaries:
        for r in s.risks:
            c = r.category.value
            if c not in cats:
                cats[c] = {"Critical": 0, "High": 0, "Medium": 0, "Low": 0, "Total": 0}
            cats[c][r.severity.value] = cats[c].get(r.severity.value, 0) + 1
            cats[c]["Total"] += 1
    if not cats:
        return
    headers = ["Category", "Critical", "High", "Medium", "Low", "Total"]
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, t in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, brand.primary_colour)
        _set_cell_margins(cell, 50, 50, 80, 80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(t)
        run.font.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for cat, counts in sorted(cats.items(), key=lambda x: -x[1]["Total"]):
        row = table.add_row()
        row.cells[0].text = cat
        for j, sev in enumerate(["Critical", "High", "Medium", "Low", "Total"], 1):
            val = counts.get(sev, 0)
            row.cells[j].text = str(val) if val > 0 else "—"
            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in row.cells:
            _set_cell_margins(cell, 40, 40, 80, 80)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    _set_table_borders(table, "D5D8DC")


def _add_footer(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:color"), "D5D8DC")
    top.set(qn("w:space"), "1")
    pBdr.append(top)
    pPr.append(pBdr)
    run = p.add_run("Generated by Portfolio Risk Copilot")
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xBD, 0xC3, 0xC7)
    run.font.italic = True


# ──────────────────────────────────────────────
# Decision generation — data-driven
# ──────────────────────────────────────────────

def _generate_decisions(report: PortfolioRiskReport, n: int = 3) -> list[str]:
    decisions: list[str] = []
    red_projects = [s for s in report.project_summaries if s.rag_status == "Red"]
    amber_projects = [s for s in report.project_summaries if s.rag_status == "Amber"]

    # Burn rate decisions — quantified
    for s in report.project_summaries:
        for r in s.risks:
            if r.category == RiskCategory.BURN_RATE and r.severity == RiskSeverity.CRITICAL and len(decisions) < n:
                decisions.append(
                    f"URGENT: {s.project_name} budget is critical — "
                    f"approve a budget top-up or cut scope before the next review cycle. "
                    f"Without action, delivery cannot be completed within allocation."
                )
                break

    # Blocked work — specific
    blocked_projects = set()
    for s in report.project_summaries:
        for r in s.risks:
            if r.category == RiskCategory.BLOCKED_WORK and r.severity in (RiskSeverity.CRITICAL, RiskSeverity.HIGH):
                blocked_projects.add(s.project_name)
    if blocked_projects and len(decisions) < n:
        names = ", ".join(list(blocked_projects)[:3])
        decisions.append(
            f"Assign resolution owners to unblock {names} — "
            f"set 5-day resolution deadlines and escalate if not cleared. "
            f"Blocked tasks are the #1 source of delivery delay this cycle."
        )

    # Red project escalation
    if red_projects and len(decisions) < n:
        names = ", ".join(s.project_name for s in red_projects[:3])
        decisions.append(
            f"Escalate {names} to executive review — "
            f"{len(red_projects)} project{'s' if len(red_projects) > 1 else ''} "
            f"at Red status with combined {sum(s.risk_count for s in red_projects)} risks. "
            f"Leadership intervention required this cycle."
        )

    # Amber watch
    if amber_projects and len(decisions) < n:
        names = ", ".join(s.project_name for s in amber_projects[:2])
        decisions.append(
            f"Monitor {names} — emerging risks could escalate to Red "
            f"without proactive mitigation. Schedule mid-cycle check-in."
        )

    # Catch-all
    while len(decisions) < n:
        decisions.append(
            "Schedule portfolio risk review in 2 weeks to track resolution "
            "progress and reassess the risk posture."
        )

    return decisions[:n]


def _generate_talking_points(report: PortfolioRiskReport) -> list[str]:
    points: list[str] = []
    total = len(report.project_summaries)
    reds = sum(1 for s in report.project_summaries if s.rag_status == "Red")
    points.append(
        f"Portfolio health: {reds} of {total} projects are Red — "
        f"this requires leadership attention, not just monitoring."
    )
    cats: dict[str, int] = {}
    for s in report.project_summaries:
        for r in s.risks:
            cats[r.category.value] = cats.get(r.category.value, 0) + 1
    if cats:
        top_cat = max(cats, key=cats.get)
        points.append(
            f"Most frequent risk type: '{top_cat}' ({cats[top_cat]} instances). "
            f"This is a systemic pattern, not a one-off — consider a targeted intervention."
        )
    points.append(
        "Key question: Are we comfortable with the current risk exposure, "
        "or do we need to reallocate resources across the portfolio?"
    )
    return points


def _generate_project_actions(summary: ProjectRiskSummary) -> list[str]:
    actions: list[str] = []
    for risk in summary.risks[:3]:
        if risk.suggested_mitigation:
            first = risk.suggested_mitigation.split(". ")[0]
            actions.append(f"{first}.")
    if not actions and summary.rag_status == "Green":
        actions.append("Continue on current trajectory. No escalation needed.")
    return actions


# ──────────────────────────────────────────────
# XML/formatting helpers
# ──────────────────────────────────────────────

def _h(brand: BrandConfig, key: str, default: str) -> str:
    return brand.custom_headings.get(key, default)


def _maybe_add_logo(doc: Document, brand: BrandConfig) -> None:
    if brand.logo_path and Path(brand.logo_path).exists():
        doc.add_picture(brand.logo_path, width=Inches(1.5))


def _set_cell_bg(cell, colour_hex: str) -> None:
    tcPr = cell._element.get_or_add_tcPr()
    shading = tcPr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tcPr.append(shading)
    shading.set(qn("w:fill"), colour_hex)
    shading.set(qn("w:val"), "clear")


def _set_cell_margins(cell, top=0, bottom=0, left=0, right=0) -> None:
    tcPr = cell._element.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    tcPr.append(margins)


def _remove_table_borders(table) -> None:
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        borders.append(el)
    tblPr.append(borders)


def _set_table_borders(table, colour: str) -> None:
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), colour)
        borders.append(el)
    tblPr.append(borders)


def _highlight_run(run, colour_hex: str) -> None:
    """Apply shading/highlight to a run (simulates badge effect)."""
    rPr = run._element.get_or_add_rPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), colour_hex)
    rPr.append(shading)


def _get_top_n_risks(report: PortfolioRiskReport, n: int = 5) -> list[Risk]:
    all_risks: list[Risk] = []
    for s in report.project_summaries:
        all_risks.extend(s.risks)
    order = {RiskSeverity.CRITICAL: 0, RiskSeverity.HIGH: 1, RiskSeverity.MEDIUM: 2, RiskSeverity.LOW: 3}
    all_risks.sort(key=lambda r: order.get(r.severity, 99))
    return all_risks[:n]
