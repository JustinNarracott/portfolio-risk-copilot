"""
Portfolio Dashboard Report generator.

A standalone 2-page executive document combining all visual
intelligence into a single at-a-glance view:

Page 1: Portfolio health dashboard (composite chart), KPI cards,
         RAG summary table
Page 2: Risk heatmap, benefits waterfall, ROI bubble chart,
         executive action summary + top 3 decisions

This is the document a CXO opens first.
"""

from __future__ import annotations

from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.artefacts.docx_generator import (
    BrandConfig, _apply_base_styles, _add_header_bar, _add_section_heading,
    _add_footer, _set_cell_bg, _set_cell_margins, _set_table_borders,
    _remove_table_borders, _add_exec_action_box, _h,
)
from src.risk_engine.engine import PortfolioRiskReport
from src.benefits.calculator import PortfolioBenefitReport
from src.investment import PortfolioInvestmentReport, InvestmentAction
from src.insights import generate_executive_summary


# KPI colour thresholds
def _kpi_colour(value: float, thresholds: tuple[float, float] = (0.7, 0.85)) -> str:
    """Return hex colour based on value vs thresholds (lower=worse)."""
    if value < thresholds[0]:
        return "E74C3C"  # Red
    elif value < thresholds[1]:
        return "F39C12"  # Amber
    return "27AE60"  # Green


def _pct_colour(pct: float, invert: bool = False) -> str:
    """Return colour for a percentage. invert=True means higher is worse (e.g. drift)."""
    if invert:
        if pct > 0.30:
            return "E74C3C"
        elif pct > 0.15:
            return "F39C12"
        return "27AE60"
    return _kpi_colour(pct)


def _add_kpi_row(doc: Document, kpis: list[tuple[str, str, str]], brand: BrandConfig) -> None:
    """Add a row of KPI cards as a table. Each KPI: (label, value, colour_hex)."""
    table = doc.add_table(rows=1, cols=len(kpis))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(table)

    for i, (label, value, colour) in enumerate(kpis):
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, "F8F9FA")
        _set_cell_margins(cell, 50, 50, 80, 80)

        # Value (large)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(colour)
        run.font.name = brand.body_font

        # Label (small)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        lab = p2.add_run(label)
        lab.font.size = Pt(7)
        lab.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        lab.font.name = brand.body_font


def _add_rag_summary_table(
    doc: Document, report: PortfolioRiskReport, projects: list | None, brand: BrandConfig,
) -> None:
    """Compact RAG summary: Project | RAG | Risks | Budget % | Status."""
    budget_map: dict[str, tuple[float, float]] = {}
    if projects:
        for p in projects:
            if p.name not in budget_map:
                budget_map[p.name] = (p.budget or 0, p.actual_spend or 0)

    headers = ["Project", "RAG", "Risks", "Budget Used", "Status"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    rag_colours = {"Red": "E74C3C", "Amber": "F39C12", "Green": "27AE60"}

    # Header row
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, brand.primary_colour)
        _set_cell_margins(cell, 30, 30, 60, 60)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = brand.body_font

    # Data rows
    sorted_summaries = sorted(
        report.project_summaries,
        key=lambda s: {"Red": 0, "Amber": 1, "Green": 2}.get(s.rag_status, 3),
    )
    for s in sorted_summaries:
        row = table.add_row()
        budget, spend = budget_map.get(s.project_name, (0, 0))
        pct_used = f"{spend / budget * 100:.0f}%" if budget > 0 else "—"

        values = [
            s.project_name[:25],
            s.rag_status,
            str(s.risk_count),
            pct_used,
            s.project_status[:15],
        ]
        for i, val in enumerate(values):
            cell = row.cells[i]
            _set_cell_margins(cell, 20, 20, 50, 50)
            # Alternating row background
            if sorted_summaries.index(s) % 2 == 1:
                _set_cell_bg(cell, "F8F9FA")
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(8)
            run.font.name = brand.body_font

            # Colour the RAG cell
            if i == 1:
                colour = rag_colours.get(val, "7F8C8D")
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(colour)

    _set_table_borders(table, "D5D8DC")


def generate_portfolio_dashboard(
    report: PortfolioRiskReport,
    benefit_report: PortfolioBenefitReport | None = None,
    investment_report: PortfolioInvestmentReport | None = None,
    projects: list | None = None,
    brand: BrandConfig | None = None,
    output_path: str | Path | None = None,
) -> Path:
    """Generate a 2-page Portfolio Dashboard Report."""
    brand = brand or BrandConfig()
    doc = Document()
    _apply_base_styles(doc, brand)

    # Tighten margins for dashboard (0.5 inch all round)
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # ── Page 1: Health Overview ──
    _add_header_bar(doc, brand, _h(brand, "dashboard_title", "Portfolio Dashboard"))

    # Executive action summary
    exec_summary = generate_executive_summary(report, benefit_report, investment_report)
    _add_exec_action_box(doc, exec_summary, brand)

    # KPI cards row 1: Portfolio health
    total = len(report.project_summaries)
    reds = sum(1 for s in report.project_summaries if s.rag_status == "Red")
    ambers = sum(1 for s in report.project_summaries if s.rag_status == "Amber")
    greens = total - reds - ambers
    total_risks = sum(s.risk_count for s in report.project_summaries)
    health_pct = greens / total if total > 0 else 0

    kpis_row1 = [
        ("PROJECTS", str(total), brand.primary_colour),
        ("RED", str(reds), "E74C3C" if reds > 0 else "27AE60"),
        ("AMBER", str(ambers), "F39C12" if ambers > 0 else "27AE60"),
        ("GREEN", str(greens), "27AE60"),
        ("TOTAL RISKS", str(total_risks), "E74C3C" if total_risks > 10 else "F39C12" if total_risks > 5 else "27AE60"),
    ]
    _add_kpi_row(doc, kpis_row1, brand)

    # KPI cards row 2: Financial & benefits
    total_budget = sum((p.budget or 0) for p in projects) if projects else 0
    total_spend = sum((p.actual_spend or 0) for p in projects) if projects else 0
    budget_pct = total_spend / total_budget if total_budget > 0 else 0

    kpis_row2 = [
        ("TOTAL BUDGET", f"£{total_budget / 1000:.0f}k", brand.primary_colour),
        ("SPENT", f"£{total_spend / 1000:.0f}k", _pct_colour(1 - budget_pct)),
        ("BUDGET USED", f"{budget_pct * 100:.0f}%", _pct_colour(budget_pct, invert=True)),
    ]

    if benefit_report:
        kpis_row2.append(
            ("BENEFITS DRIFT", f"{benefit_report.portfolio_drift_pct * 100:.0f}%",
             _pct_colour(benefit_report.portfolio_drift_pct, invert=True))
        )
    if investment_report:
        kpis_row2.append(
            ("PORTFOLIO ROI", f"{investment_report.portfolio_roi * 100:.0f}%",
             _kpi_colour(investment_report.portfolio_roi / 2))
        )

    _add_kpi_row(doc, kpis_row2, brand)

    # Dashboard composite chart (compact for tight page fit)
    try:
        from src.charts import chart_portfolio_dashboard_compact
        chart_path = chart_portfolio_dashboard_compact(report, benefit_report, investment_report, projects)
        pic = doc.add_picture(str(chart_path), width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        try:
            from src.charts import chart_portfolio_dashboard
            chart_path = chart_portfolio_dashboard(report, benefit_report, investment_report, projects)
            doc.add_picture(str(chart_path), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    # ── Page 2: Detail view ──
    doc.add_page_break()

    # RAG summary table
    _add_section_heading(doc, brand, "Project RAG Summary")
    _add_rag_summary_table(doc, report, projects, brand)

    # Side-by-side: Benefits waterfall + ROI bubble
    charts_added = False
    try:
        from src.charts import chart_benefits_waterfall, chart_roi_vs_risk
        ct = doc.add_table(rows=1, cols=2)
        ct.alignment = WD_TABLE_ALIGNMENT.LEFT
        _remove_table_borders(ct)

        if benefit_report:
            wf = chart_benefits_waterfall(benefit_report)
            ct.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ct.rows[0].cells[0].paragraphs[0].add_run().add_picture(str(wf), width=Inches(3))

        if investment_report:
            roi = chart_roi_vs_risk(investment_report)
            ct.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ct.rows[0].cells[1].paragraphs[0].add_run().add_picture(str(roi), width=Inches(3.2))

        charts_added = True
    except Exception:
        pass

    if charts_added:
        doc.add_paragraph()

    # Top 3 decisions
    _add_section_heading(doc, brand, "Key Decisions This Cycle")

    # Pull from investment recommendations
    decisions: list[str] = []
    if investment_report and investment_report.recommendations:
        decisions.extend(investment_report.recommendations[:2])
    if benefit_report and benefit_report.recommendations:
        decisions.extend(benefit_report.recommendations[:2])

    # Fallback: derive from risk data
    if not decisions:
        for s in report.project_summaries:
            if s.rag_status == "Red":
                decisions.append(f"Escalate {s.project_name} — {s.risk_count} risks at Red status.")
                if len(decisions) >= 3:
                    break

    for i, d in enumerate(decisions[:4], 1):
        p = doc.add_paragraph()
        num = p.add_run(f"  {i}  ")
        num.font.size = Pt(9)
        num.font.bold = True
        num.font.color.rgb = RGBColor.from_string(brand.accent_colour)
        text = p.add_run(d)
        text.font.size = Pt(9)
        text.font.name = brand.body_font

    _add_footer(doc)

    path = Path(output_path) if output_path else Path("portfolio-dashboard.docx")
    doc.save(str(path))
    return path
